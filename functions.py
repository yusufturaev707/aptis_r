import pandas as pd
import json
import datetime
import os
import docx
from dateutil.relativedelta import relativedelta
from docxcompose.composer import Composer
from docx import Document as Document_compose
import PyPDF2

sana = ''


def split_pdf(file, date):
    path = f"results_pdf/{date}"
    path_id = f"results_id/{date}"
    if not os.path.isdir(path):
        os.mkdir(path=path)
    if not os.path.isdir(path_id):
        os.mkdir(path=path_id)
    pdf = PyPDF2.PdfFileReader(open(file, "rb"))
    n = pdf.numPages
    k = 0
    for i in range(n):
        if i % 2 == 0:
            newpdf = PyPDF2.PdfFileWriter()
            newpdf.addPage(pdf.getPage(i))
            newpdf.addPage(pdf.getPage(1))
            k += 1
            with open(f"results_pdf/{date}/page-{k}.pdf", "wb") as f:
                newpdf.write(f)

            with open(f"results_id/{date}/page-{k}.pdf", "wb") as f:
                newpdf.write(f)


def get_all_names(file):
    xls = pd.ExcelFile(file)
    sheets = xls.sheet_names
    n = len(sheets)
    names = list()
    for i in range(0, n, 2):
        data = pd.read_excel(xls, sheet_name=sheets[i])
        df = pd.DataFrame(data)
        df.fillna(0, inplace=True)
        names.append(df.iloc[5, 3])
    index = None
    for name in names:
        if names.count(name) > 1:
            index = names.index(name)
            names[index] = f" {names[index]} "
            break

    return names


def load_id(file):
    xls = pd.ExcelFile(file)
    sheets = xls.sheet_names

    data = pd.read_excel(xls, sheet_name=sheets[0])
    df = pd.DataFrame(data)
    df.fillna(0, inplace=True)
    n = df.shape[0]
    users = []
    for i in range(n):
        id = df.iloc[i, 0]
        fam = df.iloc[i, 1]
        ism = df.iloc[i, 2]
        sharf = df.iloc[i, 3]
        seriya = df.iloc[i, 4]
        raqam = df.iloc[i, 5]
        phone = df.iloc[i, 6]
        test_date = df.iloc[i, 7]

        user = {
            "id": id,
            "fio": f"{ism} {fam}",
            "full_name": f"{fam} {ism} {sharf}",
            "passport": f"{seriya} {raqam}",
            "phone": phone,
            "test_date": test_date,
        }
        users.append(user)

    return users


def rename_file(names, date, names_id=None, skills=None, can_ref=None):
    k = 0
    users = []
    for i in range(len(names)):
        os.rename(f'results_pdf/{date}/page-{i + 1}.pdf', f'results_pdf/{date}/{names[i]}.pdf')
        id = None
        is_have = 0
        user = []
        for j in range(len(names_id)):
            if names[i] == names_id[j]['fio']:
                is_have = 1
                k += 1
                id = names_id[j]['id']
                user.append(id)
                user.append(skills[i])
                users.append(user)
                os.rename(f'results_id/{date}/page-{k}.pdf', f'results_id/{date}/{id}.pdf')
                break
        if is_have == 0:
            k += 1

    df = pd.DataFrame(users, columns=["abitur_id", "Result"])
    df.to_excel(f'results_id/{date}/res_json.xlsx')

    persons = []
    for i in range(len(can_ref)):
        person = []
        for j in range(len(names_id)):
            if can_ref[i][1] == names_id[j]['fio']:
                number = can_ref[i][0]
                full_name = names_id[j]['full_name']
                s_code = names_id[j]['id']
                passport = names_id[j]['passport']
                language = "Ingliz tili"
                regnum = can_ref[i][0]

                day_now = names_id[j]['test_date']
                day, month, year = str(day_now).split('.')
                test_date = datetime.date(int(year), int(month), int(day))
                untill_years = test_date + relativedelta(years=3)
                untill_years = untill_years.strftime("%d.%m.%y")

                term = f"{names_id[j]['test_date']}-{untill_years}"

                person.append(number)
                person.append(full_name)
                person.append(s_code)
                person.append(passport)
                person.append(language)
                person.append(regnum)
                person.append(term)
                persons.append(person)
                break
    df1 = pd.DataFrame(persons, columns=["number", "full_name", "s_code", "passport", "language", "regnum", "term"])
    df1.to_excel(f'result_excel/{date}.xlsx')


def get_month(date):
    switcher = {
        1: "yanvar",
        2: "fevral",
        3: "mart",
        4: "aprel",
        5: "may",
        6: "iyun",
        7: "iyul",
        8: "avgust",
        9: "sentabr",
        10: "oktabr",
        11: "noyabr",
        12: "dekabr",
    }
    day = date.split('.')[0]
    month = date.split('.')[1]
    return f"{int(day)}-{switcher.get(int(month), 'Noaniq')}"


def merge_fio(firstname, lastname):
    fio = f"{firstname} {lastname}"
    return fio


def get_CEFR(df, kun=None, oy=None, yil=None, value=None, start=5, step=6):
    A = []
    n = df.shape[0]

    if value != None:
        for i in range(start, n, step):
            a = []

            if df.iloc[i, 2] == value:
                a.append(merge_fio(df.iloc[i, 0], df.iloc[i, 1]))
                a.append(df.iloc[i, 2])
                A.append(a)

    if value == None:
        for i in range(5, n, 6):
            a = []
            ism = fam = gram = None

            ism = df.iloc[i, 0]
            fam = df.iloc[i, 1]
            cefr = df.iloc[i, 2]

            gram = df.iloc[i + 1, 3]
            listining = df.iloc[i + 2, 2]
            reading = df.iloc[i + 3, 2]
            speaking = df.iloc[i + 4, 2]
            writing = df.iloc[i + 5, 2]

            a.append(fam)
            a.append(ism)
            a.append(int(kun))
            a.append(int(oy))
            a.append(int(yil))
            a.append(gram)
            a.append(listining)
            a.append(reading)
            a.append(speaking)
            a.append(writing)
            a.append(cefr)
            print(a)
            A.append(a)

    return A


def load_data(file, date):
    data = pd.read_excel(file)
    df = pd.DataFrame(data, columns=['Unnamed: 3', 'Unnamed: 4', 'Unnamed: 10', 'Unnamed: 11', 'Unnamed: 17'])

    # date = df.iloc[6, 4]
    kun, oy, yil = date.split('.')
    global sana
    sana = f"{kun}.{oy}.{yil}"

    B = get_CEFR(df, value="B2")
    C = get_CEFR(df, value="C")
    All = get_CEFR(df, kun, oy, yil)

    df = pd.DataFrame(B, columns=["FIO", "CEFR"])
    df.to_excel(f'cefrs_B2/{sana}_{len(B)}.xlsx')

    df = pd.DataFrame(C, columns=["FIO", "CEFR"])
    df.to_excel(f'cefrs_C/{sana}_{len(C)}.xlsx')

    return C, B, All


def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def writetxt(data, filename):
    doc = docx.Document(filename)
    for row in data:
        doc.paragraphs[15].runs[4].text = get_month(sana)
        doc.paragraphs[15].runs[14].text = f"{row[0]}"

        path = f"xatlar/{sana}"
        if not os.path.isdir(path):
            os.mkdir(path=path)
        doc.save(f"{path}/{row[0]}.docx")


def combine_all_docx(filename_master, files_list):
    number_of_sections = len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document_compose(f'xatlar/{sana}/{files_list[i]}')
        composer.append(doc_temp)
    composer.save(f"xatlar/{sana}/{sana}.docx")
